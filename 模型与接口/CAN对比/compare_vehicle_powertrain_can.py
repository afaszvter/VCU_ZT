from __future__ import annotations

import csv
import os
import re
import subprocess
import sys
import tempfile
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

import openpyxl
from docx import Document


ROOT = Path(r"D:\VCU\VCU_SL_ZhengZhou\ZT2")
BASE_XLS = ROOT / "大台架材料" / "整车与动力信号" / "整车CAN20260630.xls"
OUTPUT_DIR = ROOT / "模型与接口" / "CAN对比"
REPORT_MD = OUTPUT_DIR / "整车与动力CAN_元件协议对比记录.md"
DETAIL_CSV = OUTPUT_DIR / "整车与动力CAN_元件协议对比明细.csv"


@dataclass
class Signal:
    name: str
    start_byte: str = ""
    start_bit: str = ""
    length: str = ""
    desc: str = ""


@dataclass
class Message:
    source_file: str
    source_kind: str
    component: str
    network: str
    msg_id: str
    msg_name: str
    cycle_ms: str = ""
    sender: str = ""
    receiver: str = ""
    compare_mode: str = "exact"
    signals: list[Signal] = field(default_factory=list)


def clean(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r", "\n").strip()
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n+", "\n", text)
    return text


def norm_id(value: str) -> str:
    text = clean(value).upper().replace(" ", "")
    text = text.replace("0X", "")
    return text


def norm_sig(value: str) -> str:
    text = clean(value).upper()
    text = re.sub(r"[\s_\-()/（）\[\]【】:：，,\.。]+", "", text)
    return text


def is_hex_id(value: str) -> bool:
    return bool(re.fullmatch(r"0[Xx][0-9A-Fa-f]+", clean(value)))


def cycle_to_number(value: str) -> str:
    text = clean(value)
    match = re.search(r"(\d+(?:\.\d+)?)", text)
    return match.group(1) if match else text


def ps_literal(path: Path | str) -> str:
    return str(path).replace("'", "''")


def read_xls_sheet(path: Path, sheet_name: str) -> list[list[str]]:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as tmp:
        csv_path = Path(tmp.name)
    ps = f"""
$ErrorActionPreference = 'Stop'
$p = '{ps_literal(path)}'
$sheet = '{sheet_name}'
$out = '{ps_literal(csv_path)}'
$conn = New-Object System.Data.OleDb.OleDbConnection("Provider=Microsoft.ACE.OLEDB.12.0;Data Source=$p;Extended Properties='Excel 8.0;HDR=NO;IMEX=1'")
$conn.Open()
$cmd = $conn.CreateCommand()
$cmd.CommandText = "SELECT * FROM [" + $sheet + "]"
$da = New-Object System.Data.OleDb.OleDbDataAdapter($cmd)
$dt = New-Object System.Data.DataTable
[void]$da.Fill($dt)
$conn.Close()
$dt | Export-Csv -LiteralPath $out -NoTypeInformation -Encoding UTF8
"""
    try:
        result = subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=True,
        )
        if result.stderr.strip():
            print(result.stderr.strip(), file=sys.stderr)
        with csv_path.open("r", encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f)
            rows: list[list[str]] = []
            if reader.fieldnames is None:
                return rows
            fieldnames = sorted(reader.fieldnames, key=lambda x: int(x[1:]) if x and x[1:].isdigit() else 9999)
            for row in reader:
                rows.append([clean(row.get(name, "")) for name in fieldnames])
            return rows
    finally:
        try:
            csv_path.unlink(missing_ok=True)
        except Exception:
            pass


def read_xlsx_sheet(path: Path, sheet_name: str) -> list[list[str]]:
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[sheet_name]
    rows: list[list[str]] = []
    for row in ws.iter_rows(values_only=True):
        rows.append([clean(cell) for cell in row])
    return rows


def parse_generic_matrix(
    rows: list[list[str]],
    source_file: str,
    source_kind: str,
    component: str,
    network: str,
    compare_mode: str = "exact",
) -> list[Message]:
    header_idx = None
    for idx, row in enumerate(rows):
        if row and "MSG NAME" in row[0].upper():
            header_idx = idx
            break
    if header_idx is None:
        return []

    messages: list[Message] = []
    current: Message | None = None
    for row in rows[header_idx + 1 :]:
        if not any(row):
            continue
        msg_name = row[0] if len(row) > 0 else ""
        msg_id = row[2] if len(row) > 2 else ""
        sig_name = row[6] if len(row) > 6 else ""
        sig_desc = row[7] if len(row) > 7 else ""
        start_byte = row[9] if len(row) > 9 else ""
        start_bit = row[10] if len(row) > 10 else ""
        bit_len = row[12] if len(row) > 12 else ""

        if is_hex_id(msg_id):
            current = Message(
                source_file=source_file,
                source_kind=source_kind,
                component=component,
                network=network,
                msg_id=norm_id(msg_id),
                msg_name=msg_name,
                cycle_ms=cycle_to_number(row[4] if len(row) > 4 else ""),
                compare_mode=compare_mode,
            )
            messages.append(current)
            continue

        if current and sig_name and "SIGNAL NAME" not in sig_name.upper():
            current.signals.append(
                Signal(
                    name=sig_name,
                    start_byte=start_byte,
                    start_bit=start_bit,
                    length=bit_len,
                    desc=sig_desc,
                )
            )
    return messages


def parse_pdu_xlsx(path: Path) -> list[Message]:
    rows = read_xlsx_sheet(path, "MultiOne")
    messages: list[Message] = []
    current: Message | None = None
    current_section = ""

    def append_pdu_signal(message: Message | None, row: list[str]) -> None:
        if message is None or len(row) <= 4 or not row[4]:
            return
        if "SIGNAL NAME" in row[4].upper():
            return
        if row[4].lower().startswith("reserve"):
            return
        message.signals.append(
            Signal(
                name=row[4],
                start_byte=row[1] if len(row) > 1 else "",
                start_bit=row[2] if len(row) > 2 else "",
                length=row[3] if len(row) > 3 else "",
                desc=row[5] if len(row) > 5 else "",
            )
        )

    for row in rows:
        first = row[0] if row else ""
        if first in {"VCU_transmit", "MultiOne_transmit"}:
            current_section = first
            continue
        if "ID =" in first:
            msg_name = clean(first.split("\n")[0])
            msg_id_match = re.search(r"ID\s*=\s*(0x[0-9A-Fa-f]+)", first)
            cycle_match = re.search(r"Cycletime\s*=\s*([0-9.]+)\s*ms", first, re.IGNORECASE)
            current = Message(
                source_file=path.name,
                source_kind="xlsx",
                component="PDU/继电器",
                network="元件协议",
                msg_id=norm_id(msg_id_match.group(1) if msg_id_match else ""),
                msg_name=msg_name,
                cycle_ms=cycle_match.group(1) if cycle_match else "",
                sender="VCU" if current_section == "VCU_transmit" else "PDU",
                receiver="PDU" if current_section == "VCU_transmit" else "VCU/BMS/ICU",
                compare_mode="exact",
            )
            messages.append(current)
            append_pdu_signal(current, row)
            continue
        append_pdu_signal(current, row)
    return messages


def parse_pmsm_docx(path: Path) -> list[Message]:
    doc = Document(path)
    summary = doc.tables[1]
    rows = summary.rows[1:]
    messages: list[Message] = []
    table_indices = {1: 2, 2: 3, 3: 4, 4: 5}
    for row in rows:
        cells = [clean(c.text) for c in row.cells]
        if not cells[0].isdigit():
            continue
        seq_no = int(cells[0])
        message = Message(
            source_file=path.name,
            source_kind="docx",
            component="永磁同步主驱",
            network="元件协议",
            msg_id=norm_id(cells[2]),
            msg_name=cells[1],
            cycle_ms=cycle_to_number(cells[4]),
            compare_mode="exact",
        )
        if seq_no in table_indices:
            sig_table = doc.tables[table_indices[seq_no]]
            for sig_row in sig_table.rows[1:]:
                sig_cells = [clean(c.text) for c in sig_row.cells]
                if not sig_cells or not sig_cells[0].isdigit():
                    continue
                message.signals.append(
                    Signal(
                        name=sig_cells[3],
                        start_bit=sig_cells[1],
                        length=sig_cells[2],
                        desc=sig_cells[8] if len(sig_cells) > 8 else "",
                    )
                )
        messages.append(message)
    return messages


def parse_aux_docx(path: Path) -> list[Message]:
    doc = Document(path)
    tables = [doc.tables[i] for i in [2, 3, 5, 7]]
    messages: list[Message] = []
    for table in tables:
        msg_id_match = re.search(r"0x[0-9A-Fa-f]+", clean(table.rows[0].cells[2].text))
        message = Message(
            source_file=path.name,
            source_kind="docx",
            component="辅驱三合一",
            network="元件协议",
            msg_id=norm_id(msg_id_match.group(0) if msg_id_match else ""),
            msg_name=clean(table.rows[0].cells[2].text),
            cycle_ms=cycle_to_number(clean(table.rows[1].cells[-1].text)),
            sender=clean(table.rows[1].cells[0].text),
            receiver=clean(table.rows[1].cells[1].text),
            compare_mode="coarse",
        )
        for row in table.rows[6:]:
            cells = [clean(c.text) for c in row.cells]
            if not cells or not cells[0].startswith("BYTE"):
                continue
            name = cells[1]
            if not name or name == "预留":
                continue
            message.signals.append(Signal(name=name, start_byte=cells[0], desc="; ".join(x for x in cells[6:] if x)))
        messages.append(message)
    return messages


def parse_brs_docx(path: Path) -> list[Message]:
    doc = Document(path)
    relevant_tables = [doc.tables[i] for i in [4, 6, 8, 9, 11, 12]]
    messages: list[Message] = []
    for table in relevant_tables:
        header = [clean(c.text) for c in table.rows[0].cells]
        msg_id_match = re.search(r"0x[0-9A-Fa-f]+", " ".join(header))
        message = Message(
            source_file=path.name,
            source_kind="docx",
            component="制动电阻器",
            network="元件协议",
            msg_id=norm_id(msg_id_match.group(0) if msg_id_match else ""),
            msg_name=clean(table.rows[0].cells[2].text),
            cycle_ms=cycle_to_number(clean(table.rows[1].cells[-1].text)),
            sender=clean(table.rows[1].cells[0].text),
            receiver=clean(table.rows[1].cells[1].text),
            compare_mode="coarse",
        )
        for row in table.rows[6:]:
            cells = [clean(c.text) for c in row.cells]
            if not cells or not cells[0].startswith("BYTE"):
                continue
            candidates = [x for x in cells[1:] if x and x not in {"预留", "/", "全写1"} and not re.fullmatch(r"Bit\d+~\d+", x, re.IGNORECASE)]
            name = candidates[0] if candidates else cells[1]
            if not name:
                continue
            message.signals.append(Signal(name=name, start_byte=cells[0]))
        messages.append(message)
    return messages


def build_total_messages() -> dict[str, list[Message]]:
    total: dict[str, list[Message]] = defaultdict(list)
    for sheet_name, network in [("整车CAN矩阵$", "整车CAN"), ("动力CAN矩阵$", "动力CAN")]:
        rows = read_xls_sheet(BASE_XLS, sheet_name)
        msgs = parse_generic_matrix(rows, BASE_XLS.name, "xls", "整车与动力总表", network, compare_mode="exact")
        for msg in msgs:
            total[msg.msg_id].append(msg)
    return total


def build_component_messages() -> list[Message]:
    messages: list[Message] = []

    prop_rows = read_xlsx_sheet(ROOT / "大台架材料" / "元件" / "三通阀" / "Φ25 V560424A15三通比例阀can协议.xlsx", "Matrix")
    messages.extend(parse_generic_matrix(prop_rows, "Φ25 V560424A15三通比例阀can协议.xlsx", "xlsx", "三通比例阀", "元件协议"))

    solenoid_rows = read_xls_sheet(ROOT / "大台架材料" / "元件" / "三通阀" / "三通电磁阀2733523 - 矩阵 - A0 - 2023.10.26.xls", "Sheet1$")
    messages.extend(parse_generic_matrix(solenoid_rows, "三通电磁阀2733523 - 矩阵 - A0 - 2023.10.26.xls", "xls", "三通电磁阀", "元件协议"))

    srm_rows = read_xls_sheet(ROOT / "大台架材料" / "元件" / "开关磁阻" / "开关磁阻电机MP2505_深蓝动力矿卡四合一通讯协议.xls", "Matrix$")
    messages.extend(parse_generic_matrix(srm_rows, "开关磁阻电机MP2505_深蓝动力矿卡四合一通讯协议.xls", "xls", "开关磁阻四合一", "元件协议"))

    messages.extend(parse_pdu_xlsx(ROOT / "大台架材料" / "元件" / "永磁电机和配电" / "PDU电机控制器扩展帧CAN通讯协议V1.0（PDU）——继电器相关.xlsx"))
    messages.extend(parse_pmsm_docx(ROOT / "大台架材料" / "元件" / "永磁电机和配电" / "扩展帧通讯协议V1.0（主驱）——只有一个电机驱动.docx"))
    messages.extend(parse_aux_docx(ROOT / "大台架材料" / "元件" / "永磁电机和配电" / "电机控制器扩展帧CAN通讯协议V1.0（辅驱三合一）——好像都是DCDC和DCAC.docx"))
    messages.extend(parse_brs_docx(ROOT / "大台架材料" / "元件" / "制动电阻" / "制动电阻器CAN通信协议.docx"))

    return messages


def choose_total_message(candidates: Iterable[Message], component: Message) -> Message | None:
    candidates = list(candidates)
    if not candidates:
        return None
    same_name = [m for m in candidates if norm_sig(m.msg_name) == norm_sig(component.msg_name)]
    if same_name:
        return same_name[0]
    return candidates[0]


def compare_messages(component_msgs: list[Message], total_map: dict[str, list[Message]]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for msg in sorted(component_msgs, key=lambda x: (x.component, x.msg_id, x.msg_name)):
        total_msg = choose_total_message(total_map.get(msg.msg_id, []), msg)
        result: dict[str, str] = {
            "component": msg.component,
            "source_file": msg.source_file,
            "msg_id": msg.msg_id,
            "component_msg_name": msg.msg_name,
            "component_cycle_ms": msg.cycle_ms,
            "component_signal_count": str(len(msg.signals)),
            "status": "",
            "total_network": "",
            "total_msg_name": "",
            "total_cycle_ms": "",
            "total_signal_count": "",
            "signal_exact_match_count": "",
            "component_only_signals": "",
            "total_only_signals": "",
            "note": "",
        }

        if total_msg is None:
            result["status"] = "总表缺失"
            result["note"] = "该元件协议中的报文ID未在整车CAN20260630总表中找到。"
            rows.append(result)
            continue

        result["total_network"] = total_msg.network
        result["total_msg_name"] = total_msg.msg_name
        result["total_cycle_ms"] = total_msg.cycle_ms
        result["total_signal_count"] = str(len(total_msg.signals))

        status_flags: list[str] = []
        if msg.cycle_ms and total_msg.cycle_ms and msg.cycle_ms != total_msg.cycle_ms:
            status_flags.append("周期不一致")

        comp_names = [s.name for s in msg.signals if s.name]
        total_names = [s.name for s in total_msg.signals if s.name]
        comp_norms = [norm_sig(x) for x in comp_names]
        total_norms = [norm_sig(x) for x in total_names]
        comp_counter = Counter(comp_norms)
        total_counter = Counter(total_norms)
        comp_name_map: dict[str, list[str]] = defaultdict(list)
        total_name_map: dict[str, list[str]] = defaultdict(list)
        for name in comp_names:
            comp_name_map[norm_sig(name)].append(name)
        for name in total_names:
            total_name_map[norm_sig(name)].append(name)
        exact_hit_count = sum(min(comp_counter[k], total_counter[k]) for k in set(comp_counter) | set(total_counter))
        result["signal_exact_match_count"] = str(exact_hit_count)

        if msg.compare_mode == "exact":
            comp_only: list[str] = []
            total_only: list[str] = []
            for key in sorted(set(comp_counter) | set(total_counter)):
                diff_comp = comp_counter[key] - total_counter[key]
                diff_total = total_counter[key] - comp_counter[key]
                if diff_comp > 0:
                    comp_only.extend(comp_name_map[key][:diff_comp])
                if diff_total > 0:
                    total_only.extend(total_name_map[key][:diff_total])
            result["component_only_signals"] = " | ".join(comp_only[:12])
            result["total_only_signals"] = " | ".join(total_only[:12])
            if comp_only or total_only:
                status_flags.append("信号不一致")
        else:
            if len(msg.signals) != len(total_msg.signals):
                status_flags.append("信号数量不一致")
            result["note"] = "元件协议为字节/中文描述式定义，已按报文ID、周期和信号数量做粗比对。"

        if not status_flags:
            result["status"] = "一致"
        else:
            result["status"] = "、".join(status_flags)

        rows.append(result)
    return rows


def write_csv(rows: list[dict[str, str]], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = [
        "component",
        "source_file",
        "msg_id",
        "component_msg_name",
        "component_cycle_ms",
        "component_signal_count",
        "status",
        "total_network",
        "total_msg_name",
        "total_cycle_ms",
        "total_signal_count",
        "signal_exact_match_count",
        "component_only_signals",
        "total_only_signals",
        "note",
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def build_report(rows: list[dict[str, str]]) -> str:
    by_component: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        by_component[row["component"]].append(row)

    total_count = len(rows)
    ok_count = sum(1 for r in rows if r["status"] == "一致")
    missing_count = sum(1 for r in rows if "总表缺失" in r["status"])
    cycle_diff_count = sum(1 for r in rows if "周期不一致" in r["status"])
    signal_diff_count = sum(1 for r in rows if "信号不一致" in r["status"] or "信号数量不一致" in r["status"])

    lines: list[str] = []
    lines.append("# 整车与动力CAN总表 vs 元件协议 对比记录")
    lines.append("")
    lines.append("## 对比范围")
    lines.append(f"- 总表基准：`{BASE_XLS}`")
    lines.append("- 总表工作表：`整车CAN矩阵`、`动力CAN矩阵`")
    lines.append("- 元件来源：三通阀、制动电阻、开关磁阻、PDU、主驱永磁、辅驱三合一协议文件")
    lines.append("")
    lines.append("## 对比方法")
    lines.append("- 先按报文 ID 建立映射，再比较报文名称、周期、信号条目。")
    lines.append("- `xls/xlsx` 矩阵类协议按消息行/信号行精确抽取。")
    lines.append("- `docx` 中字节描述式协议优先比较报文 ID、周期和信号数量；如命名口径不同，不直接判为错误。")
    lines.append("")
    lines.append("## 总体结论")
    lines.append(f"- 元件协议报文总数：`{total_count}`")
    lines.append(f"- 完全一致：`{ok_count}`")
    lines.append(f"- 总表缺失：`{missing_count}`")
    lines.append(f"- 周期不一致：`{cycle_diff_count}`")
    lines.append(f"- 信号层差异：`{signal_diff_count}`")
    lines.append("")
    lines.append("## 分部件结果")
    for component, items in sorted(by_component.items()):
        ok = sum(1 for r in items if r["status"] == "一致")
        missing = sum(1 for r in items if "总表缺失" in r["status"])
        cycle = sum(1 for r in items if "周期不一致" in r["status"])
        sig = sum(1 for r in items if "信号不一致" in r["status"] or "信号数量不一致" in r["status"])
        lines.append(f"### {component}")
        lines.append(f"- 报文数：`{len(items)}`，一致：`{ok}`，缺失：`{missing}`，周期差异：`{cycle}`，信号差异：`{sig}`")
        lines.append("")
        lines.append("| 报文ID | 元件报文名 | 总表网络 | 总表报文名 | 状态 | 备注 |")
        lines.append("| --- | --- | --- | --- | --- | --- |")
        for row in items:
            note = row["note"]
            if row["component_only_signals"] or row["total_only_signals"]:
                diff_note = []
                if row["component_only_signals"]:
                    diff_note.append(f"元件独有: {row['component_only_signals']}")
                if row["total_only_signals"]:
                    diff_note.append(f"总表独有: {row['total_only_signals']}")
                note = "；".join(x for x in [note] + diff_note if x)
            lines.append(
                f"| `{row['msg_id']}` | {row['component_msg_name']} | {row['total_network'] or '-'} | "
                f"{row['total_msg_name'] or '-'} | {row['status']} | {note or '-'} |"
            )
        lines.append("")

    lines.append("## 建议处理顺序")
    lines.append("1. 先处理“总表缺失”的报文，确认是总表漏项，还是元件协议属于未采用方案。")
    lines.append("2. 再处理“周期不一致”的报文，统一到你最终要建库的 DBC 周期口径。")
    lines.append("3. 对“信号不一致”的报文逐条核对位定义、缩放因子、偏移和枚举。")
    lines.append("4. 确认后的总表版本再作为 CANA/CANB DBC 与 Simulink I/O 模型的唯一输入基线。")
    lines.append("")
    lines.append("## 明细文件")
    lines.append(f"- 结构化明细：`{DETAIL_CSV}`")
    return "\n".join(lines) + "\n"


def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    total_map = build_total_messages()
    component_msgs = build_component_messages()
    rows = compare_messages(component_msgs, total_map)
    write_csv(rows, DETAIL_CSV)
    REPORT_MD.write_text(build_report(rows), encoding="utf-8")
    print(f"report: {REPORT_MD}")
    print(f"detail: {DETAIL_CSV}")
    print(f"messages_compared: {len(rows)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
